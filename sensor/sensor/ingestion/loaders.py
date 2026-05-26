"""Chargement et extraction de texte propre par type de fichier.

Chaque loader renvoie un `ContenuCharge` : le texte extrait (en conservant les
titres de section sous forme de lignes Markdown `#` quand c'est possible) ainsi
que le type de document et un titre par défaut. La structure ainsi préservée
est exploitée par le chunker pour rattacher chaque chunk à sa section.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from loguru import logger

from sensor.models import TypeDocument

#: Extensions reconnues -> type de document.
EXTENSIONS = {
    ".pdf": TypeDocument.PDF,
    ".docx": TypeDocument.DOCX,
    ".txt": TypeDocument.TXT,
    ".md": TypeDocument.MD,
    ".markdown": TypeDocument.MD,
    ".html": TypeDocument.HTML,
    ".htm": TypeDocument.HTML,
}


@dataclass
class ContenuCharge:
    """Résultat de l'extraction d'un fichier."""

    titre: str
    type: TypeDocument
    texte: str


def type_supporte(chemin: Path) -> bool:
    return chemin.suffix.lower() in EXTENSIONS


def charger_fichier(chemin: Path) -> ContenuCharge:
    """Charge un fichier en détectant son type via l'extension."""
    chemin = Path(chemin)
    if not chemin.exists():
        raise FileNotFoundError(f"Fichier introuvable : {chemin}")
    suffixe = chemin.suffix.lower()
    type_doc = EXTENSIONS.get(suffixe)
    if type_doc is None:
        raise ValueError(
            f"Type de fichier non supporté : {suffixe!r}. "
            f"Supportés : {', '.join(sorted(EXTENSIONS))}."
        )

    if type_doc == TypeDocument.PDF:
        texte = _charger_pdf(chemin)
    elif type_doc == TypeDocument.DOCX:
        texte = _charger_docx(chemin)
    elif type_doc == TypeDocument.HTML:
        texte = extraire_texte_html(chemin.read_text(encoding="utf-8", errors="ignore"))
    else:  # TXT et MD : lecture directe (le MD garde déjà ses '#').
        texte = chemin.read_text(encoding="utf-8", errors="ignore")

    texte = _nettoyer(texte)
    if not texte.strip():
        raise ValueError(f"Aucun texte exploitable extrait de {chemin.name}.")
    return ContenuCharge(titre=chemin.stem, type=type_doc, texte=texte)


def _charger_pdf(chemin: Path) -> str:
    from pypdf import PdfReader

    lecteur = PdfReader(str(chemin))
    morceaux: list[str] = []
    for i, page in enumerate(lecteur.pages, start=1):
        contenu = page.extract_text() or ""
        if contenu.strip():
            # On marque la page comme section pour la traçabilité.
            morceaux.append(f"# Page {i}\n{contenu}")
    if not morceaux:
        logger.warning("PDF sans texte extractible (probablement scanné) : {}", chemin.name)
    return "\n\n".join(morceaux)


def _charger_docx(chemin: Path) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(str(chemin))
    lignes: list[str] = []
    for para in doc.paragraphs:
        texte = para.text.strip()
        if not texte:
            continue
        style = (para.style.name or "").lower() if para.style else ""
        if style.startswith("heading") or style.startswith("titre"):
            lignes.append(f"# {texte}")
        else:
            lignes.append(texte)
    # Les tableaux : on aplatit en lignes "cellule | cellule".
    for table in doc.tables:
        for row in table.rows:
            cellules = [c.text.strip() for c in row.cells if c.text.strip()]
            if cellules:
                lignes.append(" | ".join(cellules))
    return "\n".join(lignes)


def extraire_texte_html(html: str) -> str:
    """Extrait le texte pertinent d'un HTML, en conservant les titres comme sections.

    Supprime les éléments non informatifs (scripts, styles, nav, footer) pour
    éviter le bruit répétitif des menus et pieds de page.
    """
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(html, "html.parser")
    for balise in soup(["script", "style", "noscript", "nav", "footer", "header", "form"]):
        balise.decompose()

    morceaux: list[str] = []
    titre_page = soup.title.string.strip() if soup.title and soup.title.string else ""
    if titre_page:
        morceaux.append(f"# {titre_page}")

    # On parcourt les éléments structurants dans l'ordre du document.
    for el in soup.find_all(["h1", "h2", "h3", "h4", "p", "li", "td", "blockquote"]):
        texte = el.get_text(" ", strip=True)
        if not texte:
            continue
        if el.name in ("h1", "h2", "h3", "h4"):
            morceaux.append(f"# {texte}")
        else:
            morceaux.append(texte)
    return "\n".join(morceaux)


def _nettoyer(texte: str) -> str:
    """Normalise les espaces et supprime les lignes vides répétées."""
    lignes = [ligne.rstrip() for ligne in texte.replace("\r\n", "\n").split("\n")]
    sortie: list[str] = []
    vide_precedente = False
    for ligne in lignes:
        if not ligne.strip():
            if not vide_precedente:
                sortie.append("")
            vide_precedente = True
        else:
            sortie.append(" ".join(ligne.split()))
            vide_precedente = False
    return "\n".join(sortie).strip()


def lister_fichiers(chemin: Path) -> list[Path]:
    """Liste récursivement les fichiers supportés d'un dossier (ou renvoie [fichier])."""
    chemin = Path(chemin)
    if chemin.is_file():
        return [chemin] if type_supporte(chemin) else []
    return sorted(p for p in chemin.rglob("*") if p.is_file() and type_supporte(p))
